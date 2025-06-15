import os
import time
from docx import Document
import google.generativeai as genai # Import the SDK


# --- Configuration ---
# Replace with your actual Gemini API Key if running locally.
# In environments like Canvas, this might be automatically handled.
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    raise ValueError("GEMINI_API_KEY environment variable not set. Please set it or replace API_KEY='' with your key.")
genai.configure(api_key=API_KEY)


# Define the range of image filenames
START_INDEX = 0 # Adjust these indices based on your image files
END_INDEX = 0 # Adjust these indices based on your image files
FILE_PREFIX = ""
FILE_SUFFIX = "" # Example: FILE_PREFIX = "image_" and FILE_SUFFIX = ".jpg"

# Directory where your images are stored (update this path)
# For example: IMAGE_DIRECTORY = "/path/to/your/images/"
IMAGE_DIRECTORY = r"" # Assuming images are in the same directory as the script for demonstration

OUTPUT_WORD_FILE = r"" # Raw string to handle backslashes

# List to store all extracted texts (optional, still used for console output)
all_extracted_texts = []

# --- Helper Functions ---

def get_image_paths(start_idx, end_idx, prefix, suffix, directory):
    """Generates a list of full image file paths."""
    image_paths = []
    for i in range(start_idx, end_idx + 1):
        filename = f"{prefix}{i:04d}{suffix}"
        full_path = os.path.join(directory, filename)
        image_paths.append(full_path)
    return image_paths

from PIL import Image

def call_gemini_vision_api_with_sdk(image_path, prompt="Extract all text from this image."):
    """
    Calls the Gemini Vision API using the official SDK to extract text from an image.
    """
    try:
        # Load the image using PIL (Pillow)
        img = Image.open(image_path)

        # Create a GenerativeModel instance
        model = genai.GenerativeModel('gemini-2.0-flash')

        # Send the prompt with the image
        response = model.generate_content([prompt, img])
        
        return response.text

    except FileNotFoundError:
        print(f"Warning: Image file not found at {image_path}. Skipping.")
        return "N/A (Image file not found)"
    except Exception as e:
        print(f"Error calling Gemini Vision API with SDK for {image_path}: {e}")
        return f"API call failed: {e}"

# --- Main Execution ---

if __name__ == "__main__":
    print("Starting image text extraction process with Gemini SDK...")
    
    document = Document()
    document.add_heading('Extracted Text from Images', level=1)
    
    image_paths_to_process = get_image_paths(START_INDEX, END_INDEX, FILE_PREFIX, FILE_SUFFIX, IMAGE_DIRECTORY)

    if not image_paths_to_process:
        print("No image paths generated. Check your start/end indices and prefixes.")
    else:
        for i, image_path in enumerate(image_paths_to_process):
            print(f"\nProcessing image {i+1}/{len(image_paths_to_process)}: {os.path.basename(image_path)}")
            
            extracted_text = call_gemini_vision_api_with_sdk(image_path)
            print(f"Extracted Text: {extracted_text}")

            all_extracted_texts.append({
                "filename": os.path.basename(image_path),
                "text": extracted_text
            })

            document.add_paragraph(f"--- File: {os.path.basename(image_path)} ---")
            document.add_paragraph(extracted_text)
            document.add_paragraph("\n")

            if i < len(image_paths_to_process) - 1:
                print(f"Waiting for 20 seconds before processing next image...")
                time.sleep(20)

    try:
        document.save(OUTPUT_WORD_FILE)
        print(f"\nAll extracted texts successfully saved to: {OUTPUT_WORD_FILE}")
    except Exception as e:
        print(f"\nError saving Word document to {OUTPUT_WORD_FILE}: {e}")
        print("Please ensure the directory exists and you have write permissions.")

    print("\n--- Console Output of All Extracted Texts ---")
    for item in all_extracted_texts:
        print(f"File: {item['filename']}\nText: {item['text']}\n{'-'*30}")

    print("\nProcess complete.")



